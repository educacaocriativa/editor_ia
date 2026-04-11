"""
Aplica controle de alterações nativo do Word (.docx / OOXML).

Cada mudança é registrada como:
- <w:del> para o texto removido
- <w:ins> para o texto inserido
com autor e data/hora, tal como o Word faz nativamente.
"""
import re
from datetime import datetime, timezone
from typing import List, Tuple, Optional
from copy import deepcopy
from lxml import etree
import docx
from docx import Document

# Namespace Word
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _get_rpr(paragraph) -> Optional[etree._Element]:
    """Copia as propriedades de formatação do primeiro run do parágrafo."""
    for run in paragraph.runs:
        rpr = run._r.find(_w("rPr"))
        if rpr is not None:
            return deepcopy(rpr)
    return None


def _make_run(text: str, rpr=None, is_del: bool = False,
              bold: bool = False) -> etree._Element:
    """Cria um elemento <w:r> com texto, com suporte a negrito."""
    r = etree.Element(_w("r"))
    # Monta rPr com negrito se necessário
    if bold or rpr is not None:
        rpr_elem = deepcopy(rpr) if rpr is not None else etree.Element(_w("rPr"))
        if bold:
            # Insere <w:b/> e <w:bCs/> no rPr
            b = etree.SubElement(rpr_elem, _w("b"))
            bcs = etree.SubElement(rpr_elem, _w("bCs"))
        r.append(rpr_elem)
    tag = _w("delText") if is_del else _w("t")
    t = etree.SubElement(r, tag)
    t.text = text
    if text != text.strip() or text.startswith(" ") or text.endswith(" "):
        t.set(f"{{{XML_NS}}}space", "preserve")
    return r


_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _segmentar_texto_bold(texto: str):
    """
    Divide o texto em segmentos (texto, bold).
    Detecta marcadores **palavra** e retorna lista de (trecho, is_bold).
    """
    segmentos = []
    ultimo = 0
    for m in _BOLD_PATTERN.finditer(texto):
        antes = texto[ultimo:m.start()]
        if antes:
            segmentos.append((antes, False))
        segmentos.append((m.group(1), True))
        ultimo = m.end()
    resto = texto[ultimo:]
    if resto:
        segmentos.append((resto, False))
    return segmentos if segmentos else [(texto, False)]


def _make_ins_bold(texto: str, author: str, date: str,
                   rev_id: int, rpr=None) -> etree._Element:
    """
    Cria <w:ins> suportando **negrito** inline via marcadores **.
    Cada segmento negritado vira um <w:r> separado com <w:b/>.
    """
    el = etree.Element(_w("ins"))
    el.set(_w("id"), str(rev_id))
    el.set(_w("author"), author)
    el.set(_w("date"), date)
    segmentos = _segmentar_texto_bold(texto)
    for trecho, is_bold in segmentos:
        el.append(_make_run(trecho, rpr, is_del=False, bold=is_bold))
    return el


def _make_del(text: str, author: str, date: str, rev_id: int, rpr=None) -> etree._Element:
    """Cria um elemento <w:del> (texto removido com controle de alterações)."""
    el = etree.Element(_w("del"))
    el.set(_w("id"), str(rev_id))
    el.set(_w("author"), author)
    el.set(_w("date"), date)
    el.append(_make_run(text, rpr, is_del=True))
    return el


def _make_ins(text: str, author: str, date: str, rev_id: int, rpr=None) -> etree._Element:
    """Cria um elemento <w:ins> (texto inserido com controle de alterações)."""
    el = etree.Element(_w("ins"))
    el.set(_w("id"), str(rev_id))
    el.set(_w("author"), author)
    el.set(_w("date"), date)
    el.append(_make_run(text, rpr, is_del=False))
    return el


def _texto_completo_paragrafo(paragraph) -> str:
    """
    Reconstrói o texto completo do parágrafo lendo todos os elementos
    <w:t> diretamente do XML — ignora fragmentação por formatação.
    """
    return "".join(
        t.text or ""
        for t in paragraph._p.iter(_w("t"))
    )


def _aplicar_mudancas_no_paragrafo(
    paragraph,
    mudancas: List[Tuple[str, str]],
    author: str,
    date: str,
    rev_id: int,
) -> int:
    """
    Aplica uma lista de (original, substituto) em um parágrafo como track changes.
    Retorna o próximo rev_id disponível.
    """
    texto_completo = _texto_completo_paragrafo(paragraph)
    rpr = _get_rpr(paragraph)
    p_elem = paragraph._p

    # Localiza posições de cada mudança
    posicoes = []
    for original, substituto in mudancas:
        idx = texto_completo.find(original)
        if idx >= 0:
            posicoes.append((idx, original, substituto))

    if not posicoes:
        return rev_id

    # Ordena por posição e remove sobreposições
    posicoes.sort(key=lambda x: x[0])
    posicoes_sem_overlap = []
    ultimo_fim = 0
    for idx, original, substituto in posicoes:
        if idx >= ultimo_fim:
            posicoes_sem_overlap.append((idx, original, substituto))
            ultimo_fim = idx + len(original)

    if not posicoes_sem_overlap:
        return rev_id

    # Remove todos os filhos do parágrafo exceto <w:pPr>
    for child in list(p_elem):
        if child.tag != _w("pPr"):
            p_elem.remove(child)

    # Reconstrói o parágrafo com os track changes
    cursor = 0
    for idx, original, substituto in posicoes_sem_overlap:
        # Texto anterior à mudança
        antes = texto_completo[cursor:idx]
        if antes:
            p_elem.append(_make_run(antes, rpr))

        # Deleção do original
        p_elem.append(_make_del(original, author, date, rev_id, rpr))
        rev_id += 1

        # Inserção do substituto (pula se for deleção pura)
        # Usa _make_ins_bold para suportar **negrito** inline
        if substituto:
            p_elem.append(_make_ins_bold(substituto, author, date, rev_id, rpr))
            rev_id += 1

        cursor = idx + len(original)

    # Texto restante após a última mudança
    restante = texto_completo[cursor:]
    if restante:
        p_elem.append(_make_run(restante, rpr))

    return rev_id


def _inserir_boxe_pos_paragrafo(
    p_elem: etree._Element,
    texto_boxe: str,
    author: str,
    date: str,
    rev_id: int,
) -> int:
    """
    Insere um novo parágrafo (Boxe Confissão de Fé) imediatamente após
    p_elem, marcado como <w:ins> no controle de alterações do Word.
    Retorna o próximo rev_id disponível.
    """
    novo_p = etree.Element(_w("p"))

    # Marca o símbolo de parágrafo (¶) como inserido
    pPr = etree.SubElement(novo_p, _w("pPr"))
    rPr_pmark = etree.SubElement(pPr, _w("rPr"))
    ins_pmark = etree.SubElement(rPr_pmark, _w("ins"))
    ins_pmark.set(_w("id"), str(rev_id))
    ins_pmark.set(_w("author"), author)
    ins_pmark.set(_w("date"), date)
    rev_id += 1

    # Conteúdo do boxe como <w:ins>
    ins_content = etree.SubElement(novo_p, _w("ins"))
    ins_content.set(_w("id"), str(rev_id))
    ins_content.set(_w("author"), author)
    ins_content.set(_w("date"), date)
    rev_id += 1

    # Run com o texto — negrito para destacar o boxe
    r = etree.SubElement(ins_content, _w("r"))
    rPr = etree.SubElement(r, _w("rPr"))
    etree.SubElement(rPr, _w("b"))
    etree.SubElement(rPr, _w("bCs"))
    t = etree.SubElement(r, _w("t"))
    t.text = texto_boxe
    t.set(f"{{{XML_NS}}}space", "preserve")

    # Insere o novo parágrafo logo após p_elem no pai
    pai = p_elem.getparent()
    if pai is not None:
        idx = list(pai).index(p_elem)
        pai.insert(idx + 1, novo_p)

    return rev_id


def aplicar_todas_as_mudancas(
    doc: Document,
    mudancas: List[dict],
    author: str = "Editor IA",
) -> None:
    """
    Ponto de entrada principal.

    Recebe o documento e uma lista de alterações no formato:
    [
        {
            "texto_original": "texto exato no documento",
            "texto_corrigido": "texto após correção",
            "tipo": "ortografia",
            "explicacao": "...",
        },
        ...
    ]
    e aplica cada alteração como controle de alterações nativo do Word.

    Tipo especial "cosmovisao_boxe": insere novo parágrafo após o parágrafo
    que contém texto_original (usa texto_boxe em vez de texto_corrigido).
    """
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    rev_id = 1

    # Separa boxes de cosmovisão das correções normais
    boxes = [
        m for m in mudancas
        if m.get("tipo") == "cosmovisao_boxe" and m.get("texto_original") and m.get("texto_boxe")
    ]
    mudancas_normais = [m for m in mudancas if m.get("tipo") != "cosmovisao_boxe"]

    # Filtra apenas mudanças reais
    mudancas_validas = [
        m for m in mudancas_normais
        if m.get("texto_original") and m.get("texto_corrigido")
        and m["texto_original"] != m["texto_corrigido"]
    ]

    def _processar_paragrafos(paragrafos):
        nonlocal rev_id
        for para in paragrafos:
            texto = para.text
            if not texto.strip():
                continue
            mudancas_do_para = [
                (m["texto_original"], m["texto_corrigido"])
                for m in mudancas_validas
                if m["texto_original"] in texto
            ]
            if mudancas_do_para:
                rev_id = _aplicar_mudancas_no_paragrafo(
                    para, mudancas_do_para, author, date_str, rev_id
                )

    # Parágrafos do corpo
    _processar_paragrafos(doc.paragraphs)

    # Parágrafos em tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                _processar_paragrafos(celula.paragraphs)

    # Inserção dos Boxes Confissão de Fé
    if not boxes:
        return

    # Coleta todos os parágrafos do corpo (inclui tabelas via XML)
    todos_p = list(doc.element.body.iter(_w("p")))
    vistos_boxes: set = set()

    for box in boxes:
        ancora = box["texto_original"]
        texto_boxe = box["texto_boxe"]

        if ancora in vistos_boxes:
            continue

        # Procura o parágrafo que contém o texto-âncora
        for p_elem in todos_p:
            texto_p = "".join(
                t.text or "" for t in p_elem.iter(_w("t"))
            )
            if ancora in texto_p:
                rev_id = _inserir_boxe_pos_paragrafo(
                    p_elem, texto_boxe, author, date_str, rev_id
                )
                vistos_boxes.add(ancora)
                break
