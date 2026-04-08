"""
Extrai conteúdo estruturado de arquivos .docx para análise.
"""
from typing import List, Dict
import docx


def extrair_paragrafos(caminho: str) -> List[Dict]:
    """
    Extrai todos os parágrafos do documento com índice e texto.
    Inclui parágrafos de tabelas.
    """
    doc = docx.Document(caminho)
    paragrafos = []

    # Parágrafos do corpo principal
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if texto:
            paragrafos.append({
                "indice": i,
                "texto": texto,
                "estilo": para.style.name,
                "fonte": "corpo",
            })

    # Parágrafos dentro de tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for para in celula.paragraphs:
                    texto = para.text.strip()
                    if texto:
                        paragrafos.append({
                            "indice": len(paragrafos),
                            "texto": texto,
                            "estilo": para.style.name,
                            "fonte": "tabela",
                        })

    return paragrafos


def montar_texto_numerado(paragrafos: List[Dict]) -> str:
    """
    Monta o texto completo com parágrafos numerados para envio ao Claude.
    Facilita a referência exata no retorno das alterações.
    """
    linhas = []
    for p in paragrafos:
        linhas.append(f"[{p['indice']}] {p['texto']}")
    return "\n\n".join(linhas)


def extrair_texto_plano(caminho: str) -> str:
    """Extrai o texto completo do documento sem numeração."""
    doc = docx.Document(caminho)
    partes = []
    for para in doc.paragraphs:
        if para.text.strip():
            partes.append(para.text)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for para in celula.paragraphs:
                    if para.text.strip():
                        partes.append(para.text)
    return "\n\n".join(partes)
