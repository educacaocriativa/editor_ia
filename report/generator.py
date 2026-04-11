"""
Gera o relatório de revisão em formato .docx.
"""
from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Paleta de cores por tipo de revisão
CORES_TIPO = {
    "ortografia":       RGBColor(0xC0, 0x39, 0x2B),   # vermelho
    "gramatica":        RGBColor(0xE7, 0x4C, 0x3C),   # laranja-vermelho
    "pontuacao":        RGBColor(0xE6, 0x7E, 0x22),   # laranja
    "crase":            RGBColor(0xD3, 0x54, 0x00),   # laranja escuro
    "coesao":           RGBColor(0x27, 0xAE, 0x60),   # verde
    "coerencia":        RGBColor(0x1E, 0x8B, 0x4C),   # verde escuro
    "estilo":           RGBColor(0x16, 0xA0, 0x85),   # verde-azulado
    "clareza":          RGBColor(0x1A, 0xBC, 0x9C),   # turquesa
    "registro":         RGBColor(0x2E, 0xCC, 0x71),   # verde claro
    "pedagogico":       RGBColor(0x29, 0x80, 0xB9),   # azul
    "vocabulario":      RGBColor(0x2E, 0x86, 0xC1),   # azul médio
    "complexidade":     RGBColor(0x21, 0x8C, 0xBE),   # azul-ciano
    "conceito":         RGBColor(0x1F, 0x61, 0x8D),   # azul escuro
    "sequencia_didatica": RGBColor(0x1A, 0x5C, 0x89), # azul muito escuro
    "tom":              RGBColor(0x82, 0x79, 0xD4),   # lilás
    "exemplo":          RGBColor(0x9B, 0x59, 0xB6),   # roxo
    "inclusao":         RGBColor(0x76, 0x44, 0x8A),   # roxo escuro
    "factual":          RGBColor(0xC0, 0x39, 0x2B),   # vermelho forte
    "factual_incerto":  RGBColor(0xF3, 0x98, 0x12),   # amarelo-laranja
    "humanizacao":      RGBColor(0x21, 0x8C, 0x74),   # verde-azulado escuro
}

NOMES_TIPO = {
    "ortografia": "Ortografia",
    "gramatica": "Gramática",
    "pontuacao": "Pontuação",
    "crase": "Crase",
    "coesao": "Coesão",
    "coerencia": "Coerência",
    "estilo": "Estilo",
    "clareza": "Clareza",
    "registro": "Registro Linguístico",
    "pedagogico": "Pedagógico",
    "vocabulario": "Vocabulário",
    "complexidade": "Complexidade Sintática",
    "conceito": "Definição de Conceito",
    "sequencia_didatica": "Sequência Didática",
    "tom": "Tom e Abordagem",
    "exemplo": "Exemplos e Analogias",
    "inclusao": "Linguagem Inclusiva",
    "factual": "Erro Factual",
    "factual_incerto": "Fato Incerto (verificar)",
    "humanizacao": "Humanização",
}


def _set_cell_bg(cell, hex_color: str):
    """Define a cor de fundo de uma célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cor_hex(cor: RGBColor) -> str:
    return f"{cor[0]:02X}{cor[1]:02X}{cor[2]:02X}"


NOMES_TIPO.update({
    "bloom_classificacao": "Bloom — Classificação das Atividades",
    "bloom_correcao": "Bloom — Correção de Atividade",
    "bloom_gabarito": "Bloom — Atualização de Gabarito/Professor",
    "bloom_progressao": "Bloom — Progressão Cognitiva",
    "bncc_codigo_incorreto": "BNCC — Código Incorreto",
    "bncc_ano_incompativel": "BNCC — Ano Incompatível",
    "bncc_componente_errado": "BNCC — Componente Errado",
    "bncc_nao_trabalhada": "BNCC — Habilidade Não Trabalhada",
    "bncc_lacuna_sugerida": "BNCC — Lacuna Sugerida",
})

CORES_TIPO.update({
    "bloom_classificacao": RGBColor(0xE6, 0x7E, 0x22),   # laranja Bloom
    "bloom_correcao":      RGBColor(0xD3, 0x54, 0x00),   # laranja escuro
    "bloom_gabarito":      RGBColor(0xA0, 0x40, 0x00),   # marrom-laranja
    "bloom_progressao":    RGBColor(0xCA, 0x6F, 0x1E),
    "bncc_codigo_incorreto": RGBColor(0x8E, 0x44, 0xAD),
    "bncc_ano_incompativel": RGBColor(0x76, 0x44, 0x8A),
    "bncc_componente_errado": RGBColor(0x6C, 0x3A, 0x83),
    "bncc_nao_trabalhada": RGBColor(0xCB, 0x4B, 0x16),
    "bncc_lacuna_sugerida": RGBColor(0x26, 0x8B, 0xD2),
    "abordagem_didatica": RGBColor(0x29, 0x80, 0xB9),
    "sequencia": RGBColor(0x1A, 0x5C, 0x89),
    "tom_registro": RGBColor(0x82, 0x79, 0xD4),
    "exemplo_analogia": RGBColor(0x9B, 0x59, 0xB6),
    "scaffolding": RGBColor(0x1F, 0x61, 0x8D),
})


# Cores dos níveis de Bloom
CORES_BLOOM_NIVEL = {
    1: "FEF9E7",  # amarelo claro — Lembrar
    2: "FDEBD0",  # laranja claro — Compreender
    3: "D5F5E3",  # verde claro — Aplicar
    4: "D6EAF8",  # azul claro — Analisar
    5: "E8DAEF",  # lilás claro — Avaliar
    6: "FADBD8",  # rosa claro — Criar
}

NOMES_BLOOM_NIVEL = {
    1: "N1 — Lembrar",
    2: "N2 — Compreender",
    3: "N3 — Aplicar",
    4: "N4 — Analisar",
    5: "N5 — Avaliar",
    6: "N6 — Criar",
}

STATUS_BLOOM_EMOJI = {
    "adequado":  "✅",
    "abaixo":    "⬇",
    "acima":     "⬆",
    "estagnado": "⚠",
}


def _gerar_secao_bloom(doc: Document, etapa_bloom: Dict) -> None:
    """Gera a seção completa de Taxonomia de Bloom no relatório."""
    classificacoes = etapa_bloom.get("classificacoes", [])
    diagnostico = etapa_bloom.get("diagnostico", {})
    n_correcoes = etapa_bloom.get("total", 0)

    doc.add_heading("Taxonomia de Bloom — Análise das Atividades", level=1)

    # ── Diagnóstico geral ────────────────────────────────────────────────────
    if diagnostico:
        doc.add_heading("Diagnóstico Geral do Capítulo", level=2)

        julgamento = diagnostico.get("julgamento", "")
        p_julg = doc.add_paragraph()
        run = p_julg.add_run(f"Julgamento: {julgamento}")
        run.bold = True
        if "BOM" in julgamento:
            run.font.color.rgb = RGBColor(0x1E, 0x8B, 0x4C)
        elif "REGULAR" in julgamento:
            run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
        else:
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        niveis_enc = diagnostico.get("niveis_encontrados", [])
        niveis_falt = diagnostico.get("niveis_faltando", [])
        doc.add_paragraph(
            f"Níveis encontrados: "
            + ", ".join(NOMES_BLOOM_NIVEL.get(n, str(n)) for n in niveis_enc)
        )
        if niveis_falt:
            p_falt = doc.add_paragraph(
                "Níveis ausentes: "
                + ", ".join(NOMES_BLOOM_NIVEL.get(n, str(n)) for n in niveis_falt)
            )
            p_falt.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        doc.add_paragraph(
            f"Total de atividades analisadas: "
            f"{diagnostico.get('total_atividades', 0)} | "
            f"Correções propostas: {n_correcoes}"
        )
        doc.add_paragraph()

    # ── Tabela de classificação ───────────────────────────────────────────────
    if classificacoes:
        doc.add_heading("Classificação Individual das Atividades", level=2)

        # ── Tabela de classificação com 4 campos de Bloom ────────────────
        tabela = doc.add_table(rows=1, cols=7)
        tabela.style = "Table Grid"
        cabecalhos = [
            "Nº", "Enunciado (resumido)", "Verbo / Nível",
            "Objetivos", "Processos", "Resultantes", "Status"
        ]
        for i, cab in enumerate(cabecalhos):
            cell = tabela.cell(0, i)
            cell.paragraphs[0].add_run(cab).bold = True
            _set_cell_bg(cell, "2C3E50")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for item in classificacoes:
            linha = tabela.add_row()
            nivel_atual = item.get("nivel_bloom", 0)
            status = item.get("status", "")
            cor_linha = CORES_BLOOM_NIVEL.get(nivel_atual, "FFFFFF")

            # Nº
            linha.cells[0].paragraphs[0].add_run(str(item.get("numero", "")))
            _set_cell_bg(linha.cells[0], cor_linha)

            # Enunciado resumido
            orig = item.get("texto_original", "")
            resumo = orig[:90] + "..." if len(orig) > 90 else orig
            linha.cells[1].paragraphs[0].add_run(resumo)
            _set_cell_bg(linha.cells[1], cor_linha)

            # Verbo / Nível
            verbo = item.get("verbo", item.get("nivel_bloom_nome", ""))
            nome_nivel = NOMES_BLOOM_NIVEL.get(nivel_atual, "")
            r_v = linha.cells[2].paragraphs[0].add_run(
                f"N{nivel_atual} — {verbo}"
            )
            r_v.bold = True
            _set_cell_bg(linha.cells[2], cor_linha)

            # Objetivos
            linha.cells[3].paragraphs[0].add_run(item.get("objetivos", ""))
            _set_cell_bg(linha.cells[3], cor_linha)

            # Processos
            linha.cells[4].paragraphs[0].add_run(item.get("processos", ""))
            _set_cell_bg(linha.cells[4], cor_linha)

            # Resultantes
            linha.cells[5].paragraphs[0].add_run(item.get("resultantes", ""))
            _set_cell_bg(linha.cells[5], cor_linha)

            # Status
            emoji = STATUS_BLOOM_EMOJI.get(status, "")
            r_st = linha.cells[6].paragraphs[0].add_run(f"{emoji} {status}")
            if status in ("abaixo", "estagnado"):
                r_st.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                r_st.bold = True

        doc.add_paragraph()

        # ── Tabela de reescritas propostas ────────────────────────────────
        correcoes_bloom = [
            e for e in (etapa_bloom.get("classificacoes", []) or [])
            if False  # classificacoes já está separado acima
        ]
        # Busca correcoes no etapa_bloom diretamente
        todas_bloom = etapa_bloom.get("_todos_itens", [])
        reescritas = [
            i for i in todas_bloom
            if i.get("tipo") == "bloom_correcao"
        ]
        if reescritas:
            doc.add_heading("Reescritas Propostas — Elevação Cognitiva", level=2)
            for rew in reescritas:
                p_num = doc.add_paragraph()
                r_num = p_num.add_run(
                    f"Atividade {rew.get('numero', '')} — "
                    f"N{rew.get('nivel_original','')} ({rew.get('verbo_original','')}) "
                    f"→ N{rew.get('nivel_novo','')} ({rew.get('verbo_novo','')})"
                )
                r_num.bold = True
                r_num.font.color.rgb = RGBColor(0xD3, 0x54, 0x00)

                doc.add_paragraph(
                    f"Original: {rew.get('texto_original','')}"
                )
                p_new = doc.add_paragraph()
                p_new.add_run("Proposta: ").bold = True
                p_new.add_run(rew.get("texto_corrigido", ""))

                campos = []
                for campo in ("objetivos_novo", "processos_novo", "resultantes_novo"):
                    v = rew.get(campo, "")
                    if v:
                        campos.append(f"{campo.replace('_novo','').title()}: {v}")
                if campos:
                    doc.add_paragraph(" | ".join(campos))

                just = rew.get("justificativa", "")
                if just:
                    p_just = doc.add_paragraph(f"Justificativa: {just}")
                    p_just.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

                doc.add_paragraph()

        # ── Atualizações de gabarito / manual do professor ────────────────
        gabarito_atualizacoes = etapa_bloom.get("_gabarito", []) or []
        if gabarito_atualizacoes:
            doc.add_heading(
                "Gabarito / Manual do Professor — Atualizações",
                level=2,
            )
            doc.add_paragraph(
                "As seguintes passagens do gabarito ou orientações ao "
                "professor foram atualizadas para refletir os novos "
                "enunciados reformulados pela Taxonomia de Bloom:"
            )
            for gab in gabarito_atualizacoes:
                p_orig = doc.add_paragraph()
                p_orig.add_run("Original: ").bold = True
                p_orig.add_run(gab.get("texto_original", ""))

                p_new = doc.add_paragraph()
                p_new.add_run("Atualizado: ").bold = True
                r_new = p_new.add_run(gab.get("texto_corrigido", ""))
                r_new.font.color.rgb = RGBColor(0xA0, 0x40, 0x00)

                exp = gab.get("explicacao", "")
                if exp:
                    p_exp = doc.add_paragraph(f"Motivo: {exp}")
                    p_exp.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

                doc.add_paragraph()


def gerar_relatorio(
    caminho_saida: str,
    nome_documento: str,
    faixa_etaria: str,
    todas_as_mudancas: List[Dict],
    etapas: List[Dict],
    perfil_nome: str = "",
) -> None:
    """Gera o relatório de revisão em .docx."""
    doc = Document()

    # ── Estilos globais ──────────────────────────────────────────────────────
    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)

    # ── Cabeçalho ────────────────────────────────────────────────────────────
    titulo = doc.add_heading("Relatório de Revisão Editorial", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Metadados
    tabela_meta = doc.add_table(rows=4, cols=2)
    tabela_meta.style = "Table Grid"
    dados_meta = [
        ("Documento", nome_documento),
        ("Data da revisão", datetime.now().strftime("%d/%m/%Y %H:%M")),
        ("Público-alvo", perfil_nome or faixa_etaria),
        ("Total de alterações", str(len(todas_as_mudancas))),
    ]
    for i, (chave, valor) in enumerate(dados_meta):
        tabela_meta.cell(i, 0).paragraphs[0].add_run(chave).bold = True
        tabela_meta.cell(i, 1).paragraphs[0].add_run(valor)

    doc.add_paragraph()

    # ── Resumo por categoria ─────────────────────────────────────────────────
    doc.add_heading("Resumo por Categoria", level=1)

    if etapas:
        tabela_resumo = doc.add_table(rows=1, cols=3)
        tabela_resumo.style = "Table Grid"
        cabecalhos = ["Categoria", "Alterações", "Observações"]
        for i, cab in enumerate(cabecalhos):
            tabela_resumo.cell(0, i).paragraphs[0].add_run(cab).bold = True
            _set_cell_bg(tabela_resumo.cell(0, i), "D5E8F5")

        for etapa in etapas:
            linha = tabela_resumo.add_row()
            linha.cells[0].paragraphs[0].add_run(
                NOMES_TIPO.get(etapa["tipo"], etapa["tipo"])
            )
            linha.cells[1].paragraphs[0].add_run(str(etapa.get("total", 0)))
            obs = ""
            if etapa.get("incertos"):
                obs = f"{etapa['incertos']} fato(s) incerto(s) — verificar manualmente"
            linha.cells[2].paragraphs[0].add_run(obs)

    doc.add_paragraph()

    # ── Seção Bloom (se houver) ──────────────────────────────────────────────
    etapa_bloom = next(
        (e for e in etapas if e.get("tipo") == "bloom"), None
    )
    if etapa_bloom and (
        etapa_bloom.get("classificacoes") or etapa_bloom.get("total", 0) > 0
    ):
        _gerar_secao_bloom(doc, etapa_bloom)

    # ── Listagem detalhada por tipo ──────────────────────────────────────────
    doc.add_heading("Detalhamento das Alterações", level=1)

    # Agrupa por tipo (excluindo tipos já com seção própria na seção Bloom)
    _TIPOS_COM_SECAO_PROPRIA = {"bloom_classificacao", "bloom_correcao", "bloom_gabarito"}
    por_tipo: Dict[str, List[Dict]] = {}
    for m in todas_as_mudancas:
        t = m.get("tipo", "outros")
        if t in _TIPOS_COM_SECAO_PROPRIA:
            continue  # já aparece na seção Bloom
        por_tipo.setdefault(t, []).append(m)

    for tipo, mudancas in por_tipo.items():
        cor = CORES_TIPO.get(tipo, RGBColor(0x00, 0x00, 0x00))
        nome_tipo = NOMES_TIPO.get(tipo, tipo.title())

        heading = doc.add_heading(f"{nome_tipo} ({len(mudancas)})", level=2)
        heading.runs[0].font.color.rgb = cor

        tabela = doc.add_table(rows=1, cols=3)
        tabela.style = "Table Grid"

        # Cabeçalho
        cabecalhos = ["Original", "Corrigido", "Explicação"]
        for i, cab in enumerate(cabecalhos):
            cell = tabela.cell(0, i)
            cell.paragraphs[0].add_run(cab).bold = True
            _set_cell_bg(cell, f"{_cor_hex(cor)}")

        for m in mudancas:
            linha = tabela.add_row()
            linha.cells[0].paragraphs[0].add_run(m.get("texto_original", ""))
            corrigido = m.get("texto_corrigido", "")
            run_cor = linha.cells[1].paragraphs[0].add_run(corrigido)
            if tipo not in ("factual_incerto",):
                run_cor.font.color.rgb = cor
            linha.cells[2].paragraphs[0].add_run(m.get("explicacao", ""))

        doc.add_paragraph()

    # ── Rodapé / instruções ──────────────────────────────────────────────────
    doc.add_heading("Instruções ao Revisor Humano", level=1)
    instrucoes = [
        "1. Abra o documento revisado no Microsoft Word ou LibreOffice.",
        "2. Acesse a aba 'Revisão' e selecione 'Mostrar Marcações' para visualizar todas as alterações.",
        "3. Aceite ou rejeite cada alteração individualmente conforme seu julgamento.",
        "4. Preste atenção especial aos itens marcados como 'Fato Incerto' — estes requerem verificação manual.",
        "5. Ao finalizar, salve o documento sem o controle de alterações ativo.",
    ]
    for instrucao in instrucoes:
        doc.add_paragraph(instrucao)

    doc.save(caminho_saida)
